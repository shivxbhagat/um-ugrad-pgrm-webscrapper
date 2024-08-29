#coded by: Shiv Bhagat

#notes: 
#1. This script is not able to extract images from the elements, need to manually add images to the doc
#2. Ordered list numbering is not resetting, need to manually reset the numbering while checking the doc
#3. Must Review all the tables generated in the doc, thee are some issues with the tables in formatting and paragraph inside the table, especially multiline texts, generated as single line and need to separate them manually
#4. Has to manually type Faculty/college/school name in the doc on the top of the program
#5. Nested Unordered is generated twice, need to remove one of them
#6. By default, "On this page" and "Application Deadlines" are added, need to remove them manually as per the requirement
#7. Sometimes, with in page links, like that of "#Section-1", its not taking text related to it into account, need to manually add them, mostly happens in top 1-3 sections with "Notes" and "Academic Requirements" sections - for sure there in section 6 3rd point. Or the section's link is in the last sentence.
#8. There is separate function for English Language Proficiency, need to run that separately

# to run the script, open terminal and run the command: 
# ~ git clone https://github.com/shivxbhagat/um-ugrad-pgrm-webscrapper.git
# ~ cd um-ugrad-pgrm-webscrapper
# ~ python3 um_ugrad_prgm_scapper.py > debug.txt
# this will save the debug logs in the debug.txt file, fing for "------------------------------------------------------------" to see if any element is skipped with text in it

import subprocess
import sys

# Function to install required packages
def install_packages(packages):
    for package in packages:
        try:
            __import__(package)
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# List of required packages
required_packages = ['requests', 'bs4', 'docx']

# Install missing packages
install_packages(required_packages)

# Import the installed packages
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os

def add_element_to_doc(doc, element):
    # Recursively add HTML elements to the Word document

    print(f"Processing element: {element.name} with class: {element.get('class')}")
    
    if element.name == 'p':

        if element.get('class') == ['notification']:
            text = "Notification(i): " + element.get_text()
            doc.add_paragraph(text)
        else: 
            #if para has link in the text then add link to the doc
            linkP = element.find('a')
            if linkP:
                doc.add_paragraph(element.get_text(), style='Normal').add_run(f" ({linkP['href']})")
            else:
                doc.add_paragraph(element.get_text(), style='Normal')

    
    elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        doc.add_heading(element.get_text(), level=int(element.name[1]))
    
    elif element.name == 'ul':
        for li in element.find_all('li'):
            link = li.find('a')
            if link:
                doc.add_paragraph(link.get_text(), style='List Bullet').add_run(f" ({link['href']})")
            else: 
                doc.add_paragraph(li.get_text(), style='List Bullet')
    
    elif element.name == 'ol':
        for li in element.find_all('li'):
            linkO = li.find('a')
            if linkO:
                doc.add_paragraph(linkO.get_text(), style='List Number').add_run(f" ({linkO['href']})")
            else:
                doc.add_paragraph(li.get_text(), style='List Number')
    
    elif element.name == 'table':

        #if there is Caption and heading in table
        caption = element.find('caption')
        if caption:
            doc.add_paragraph(caption.get_text())

        # Add table to the document

        print("Adding table")
        rows = element.find_all('tr')
        cols = rows[0].find_all(['th', 'td'])
        table = doc.add_table(rows=len(rows), cols=len(cols))

        # Apply border of 1px solid black to the table
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        table.width = Inches(6.0)

        for i, row in enumerate(rows):
            row_cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(row_cells):
                cell_text = cell.get_text(strip=True)
                table_cell = table.cell(i, j)
                table_cell.text = cell_text

                # Apply bold to header cells
                if cell.name == 'th':
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        #add break after table
        doc.add_paragraph('')

    elif element.name == 'img':
        img_url = element['src']
        print(f"Adding image from URL: {img_url}")
        if img_url.startswith('http'):
            img_response = requests.get(img_url)
            img_path = os.path.join('temp_image.png')
            with open(img_path, 'wb') as img_file:
                img_file.write(img_response.content)
            doc.add_picture(img_path, width=Inches(4.0))
            os.remove(img_path)
        else:
            #local url image
            doc.add_picture(img_url, width=Inches(4.0))
    
    elif element.name == 'a':
        print(f"Adding link: {element.get_text()} ({element['href']})")
        doc.add_paragraph(f"Link: {element.get_text()} ({element['href']})")
    
    elif element.name == 'div' or element.name == 'section' or element.name == 'article' or element.name == 'main':
        print(f"Processing div/section/article element")
        for child in element.children:
            if child.name:
                add_element_to_doc(doc, child)  # Recursively process children of the div

    else: 
        print(f"Skipping element: {element.name} with class: {element.get('class')} and text: {element.get_text()}")
        print("------------------------------------------------------------")
        # Skip other elements for now


def process_url(url):
    # Fetch and parse HTML content for a single URL
    response = requests.get(url)
    html_content = response.content
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Create a new Word document for this URL
    doc = Document()

    # Change the page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Process heading
    heading = soup.find('h1', class_='page-header__heading')
    if heading:
        doc.add_heading(heading.get_text().strip(), level=1).paragraph_format.alignment = 1
    headingText = heading.get_text().strip()

    #remove (2024-2025) from the heading
    headingText = headingText.split('(')[0]

    #remove special characters from the heading
    headingText = headingText.replace(":", "")
    headingText = headingText.replace("?", "")
    headingText = headingText.replace("/", "")
    headingText = headingText.replace("\\", "")
    headingText = headingText.replace("*", "")
    headingText = headingText.replace("<", "")
    headingText = headingText.replace(">", "")
    headingText = headingText.replace("|", "")
    headingText = headingText.replace("-", "")



    # Process specific section
    section = soup.find('div', class_='js-hero-container')
    if section:
        add_element_to_doc(doc, section)
    else:
        print(f"-------- Section with class 'js-hero-container' not found in {url} -------- ")

    # doc_name
    doc_name = f"{headingText}.docx"

    # Save the Word document
    doc.save(doc_name)

    print(f"Document saved as: {doc_name}")


urlsTest = [
    #test urls
   
]

# List of URLs to process
urls = [
    #all urls
    #faculty of kinesiology have same specifications for all its programs
    #nothing there for ELC
    #IEEQ is suspended engineering program
    #BPRN is suspended nursing program
    #BSW - sociology and fort garry BSW - social work have same specifications

    #Direct Entry - U1
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/direct-entry",

    #Advanced Entry Programs
    #Faculty of Agricultural and Food Sciences
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/faculty-agricultural-and-food-sciences",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/agriculture-diploma",

    #Faculty of Architecture(Environmental Design)
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/environmental-design",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/ampp",

    #School of Art
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/fine-arts",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/fine-art-history",

    #Faculty of Arts
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/arts",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/art-is",

    #Asper School of Business
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/business-track",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/asper-school-business-track",

    #School of dental hygiene
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/dental-hygiene",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/dental-hygiene-degree",

    #Dr. Gerald Niznick College of Dentistry
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/dentistry",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/iddp",

    #Faculty of Education
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/education",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/pbde",

    #Faculty of Engineering
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/price-faculty-engineering",

    #Faculty of Environment, Earth, and Resources
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/riddell",

    #Extended Education
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/general-studies",

    #Interdisciplinary Health Programs
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/health-science-studies",

    #Faculty of Kinesiology and Recreation Management
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/kin-rec",

    #Faculty of Law
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/law",

    #Max Rady College of Medicine
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/medicine",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/pbd-medical-phys",

    #Faculty of Music
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/music",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/post-bacc-diploma-performance",

    #College of Nursing
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/nursing",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/admissions-ucn-nursing",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/midwifery",

    #College of Pharmacy
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/pharmacy",

    #College of Rehabilitation Sciences
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/respiratory-therapy",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/respiratory-therapy-dc",

    #Faculty of Science
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/science",

    #Faculty of Social Work
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/social-work-distance",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/social-work",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/social-work-inner-city-access",
    "https://umanitoba.ca/explore/undergraduate-admissions/requirements/social-work-northern-access"

    #can add more urls
]

#for English Language Proficiency
elpUrl = "https://umanitoba.ca/admissions/undergraduate/requirements/english-language-proficiency"
def elpExtract (elpUrl):
    response = requests.get(elpUrl)
    html_content = response.content
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Create a new Word document for this URL
    doc = Document()

    # Change the page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)



    # Process specific section
    section = soup.find('div', class_='content')
    if section:
        add_element_to_doc(doc, section)
    else:
        print(f"-------- Section with class 'content' not found in {elpUrl} -------- ")

    # doc_name
    doc_name = f"English_Language_Proficiency.docx"

    # Save the Word document
    doc.save(doc_name)

    print(f"Document saved as: {doc_name}")

# Process English Language Proficiency
elpExtract(elpUrl)


# Process each URL
for url in urlsTest:
    print(f"Processing URL: {url}")
    process_url(url)

print("All URLs successfully processed and saved to Word documents.")


